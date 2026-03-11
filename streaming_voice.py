import sys
import threading
import time
import os
import datetime
import subprocess

# 🌟 核心防闪退机制 1：禁止底层 Tokenizer 的多进程 fork 行为。
os.environ["TOKENIZERS_PARALLELISM"] = "false"

import numpy as np
import sounddevice as sd
import soundfile as sf
import mlx_whisper
import docx

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QTextEdit, QVBoxLayout, QHBoxLayout, 
    QWidget, QPushButton, QLabel, QFileDialog, QMessageBox
)
from PyQt6.QtCore import Qt, pyqtSignal, QObject, QTimer
from PyQt6.QtGui import QKeySequence, QShortcut

# --------- 配置区域 ---------
MODEL = "mlx-community/whisper-small-mlx"
SAMPLE_RATE = 16000
CHANNELS = 1
# --------------------------

class Signals(QObject):
    update_text = pyqtSignal(str, bool)
    status_update = pyqtSignal(str, str)
    btn_update = pyqtSignal(str, bool)
    
    # 专门用于长音频导出任务的信号
    export_status = pyqtSignal(str, str)

class VoiceInputApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("语音捕捉与转录中心")
        self.resize(360, 260)
        # 固定悬浮最前
        self.setWindowFlags(Qt.WindowType.WindowStaysOnTopHint)
        
        # 居中显示
        screen = QApplication.primaryScreen().geometry()
        x = (screen.width() - self.width()) // 2
        y = (screen.height() - self.height()) // 2
        self.move(x, y)

        self.signals = Signals()
        self.signals.update_text.connect(self.on_update_text)
        self.signals.status_update.connect(self.on_status_update)
        self.signals.btn_update.connect(self.on_btn_update)
        self.signals.export_status.connect(self.on_export_status)

        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)

        # ====== UI 构建 ====== 
        # 状态提示
        self.status_label = QLabel("✅ 初始化完毕，等待录制或导入")
        self.status_label.setStyleSheet("font-size: 12px; color: green;")
        layout.addWidget(self.status_label, alignment=Qt.AlignmentFlag.AlignCenter)

        # 连续草稿本 (文本框)
        self.text_area = QTextEdit()
        self.text_area.setReadOnly(True)
        self.text_area.setStyleSheet("background-color: #282c34; color: #abb2bf; font-size: 14px; padding: 5px;")
        layout.addWidget(self.text_area)
        
        # ====== 按钮布局 (主操作区域) ======
        btn_layout = QHBoxLayout()
        
        self.record_btn = QPushButton("🔴 开始录制 (Command+R)")
        self.record_btn.setStyleSheet("font-size: 13px; padding: 8px; background-color: #e0e0e0; color: black; border-radius: 5px;")
        self.record_btn.clicked.connect(self.toggle_recording)
        btn_layout.addWidget(self.record_btn, stretch=2)

        self.clear_btn = QPushButton("🗑️ 清空")
        self.clear_btn.setStyleSheet("font-size: 13px; padding: 8px; background-color: #ffcccc; color: black; border-radius: 5px;")
        self.clear_btn.clicked.connect(self.clear_notebook)
        btn_layout.addWidget(self.clear_btn, stretch=1)
        
        layout.addLayout(btn_layout)

        # ====== 长图文导入功能 ======
        self.import_btn = QPushButton("📁 导入音频并转为带时间戳 Word")
        self.import_btn.setStyleSheet("font-size: 13px; padding: 8px; background-color: #cce5ff; color: black; border-radius: 5px;")
        self.import_btn.clicked.connect(self.import_audio_action)
        layout.addWidget(self.import_btn)

        # ====== 内部状态管理 ======
        self.is_recording = False
        self.audio_data = []
        self.recording_thread = None
        self.transcribe_thread = None
        self.temp_file = "stream_temp.wav"
        self.mlx_lock = threading.Lock()
        self.audio_lock = threading.Lock() # 🔥 音频数据读写锁
        
        # 🌟 “连续草稿本”的核心：记忆变量
        self.committed_text = ""
        
        # 快捷键机制
        self.shortcut = QShortcut(QKeySequence("Ctrl+R"), self)
        self.shortcut.activated.connect(self.toggle_recording)

    # ---------------- 核心草稿本机制 ----------------
    def toggle_recording(self):
        if not self.is_recording:
            self.start_recording()
            subprocess.Popen(["afplay", "/System/Library/Sounds/Ping.aiff"])
        else:
            self.stop_recording()

    def start_recording(self):
        self.is_recording = True
        self.signals.btn_update.emit("⏹️ 结束本段 (追加并复制)", True)
        self.signals.status_update.emit("🎙️ 正在倾听，请说话... (可随时停顿)", "red")
        
        self.audio_data = [] # 每次只清空当前段落音频，千万别清空文本！
        
        self.recording_thread = threading.Thread(target=self.record_loop, daemon=True)
        self.recording_thread.start()
        
        self.transcribe_thread = threading.Thread(target=self.transcribe_loop, daemon=True)
        self.transcribe_thread.start()
        
    def stop_recording(self):
        self.is_recording = False
        self.signals.btn_update.emit("⏳ 正在定稿 (在此期间可随意操作界面)...", False)
        self.signals.status_update.emit("🔄 后台正在为您生成并拼接最终文本...", "blue")
        
        # 【外科手术二：UI 物理隔离】将死锁计算流放进清道夫后台线程
        threading.Thread(target=self._finalize_recording, daemon=True).start()
        
    def _finalize_recording(self):
        print(f"[{datetime.datetime.now().strftime('%H:%M:%S')}] 🏁 开始定稿流程...")
        try:
            if self.recording_thread:
                # 给 2 秒超时，防止音频驱动挂起导致死等
                self.recording_thread.join(timeout=2.0)
                
            # 进行最后一次定版转写，并打上 final=True 盖章戳
            self.do_transcribe(final=True)
        except Exception as e:
            print(f"❌ 定稿后台发生异常: {e}")
        finally:
            self.signals.btn_update.emit("🔴 继续下一段录制 (Command+R)", True)
            self.signals.status_update.emit("✅ 单段录音落锤，已无缝追加并复制", "green")
            print(f"[{datetime.datetime.now().strftime('%H:%M:%S')}] ✨ 全流程处理完毕，界面已恢复。")
        
    def record_loop(self):
        def callback(indata, frames, time, status):
            if status: pass
            with self.audio_lock:
                self.audio_data.append(indata.copy())
            
        with sd.InputStream(samplerate=SAMPLE_RATE, channels=CHANNELS, callback=callback):
            while self.is_recording:
                time.sleep(0.1)
                
    def transcribe_loop(self):
        while self.is_recording:
            time.sleep(1.5)
            if len(self.audio_data) > 0 and self.is_recording:
                self.do_transcribe(final=False)
                
    def do_transcribe(self, final=False):
        # 只有在 final 时才打印耗时，避免实时更新时刷屏
        if final: print(f"[{datetime.datetime.now().strftime('%H:%M:%S')}] 正在调用 GPU 模型...")

        with self.audio_lock:
            if not self.audio_data: return
            audio_np = np.concatenate(self.audio_data, axis=0)
        
        try:
            # 🔥 策略二：VAD 物理门禁
            if np.max(np.abs(audio_np)) < 0.005:
                return
            
            # 🌟 优化：直接传递内存数组给 mlx_whisper，跳过磁盘 WAV 写入进度，彻底避免文件死锁
            with self.mlx_lock:
                result = mlx_whisper.transcribe(
                    audio_np, 
                    path_or_hf_repo=MODEL,
                    language="zh"
                )
            text = result.get('text', '').strip()
            if text:
                self.signals.update_text.emit(text, final)
            if final: print(f"[{datetime.datetime.now().strftime('%H:%M:%S')}] 模型转录成功。")
        except Exception as e:
            if final: print(f"⚠️ 转录过程失败: {e}")
            pass

    def on_update_text(self, text, final):
        # 草稿本的核心展示逻辑：已经确定的最终文字 + 当前这一小段试探性的文字
        separator = "\n\n" if self.committed_text else ""
        display_text = self.committed_text + separator + text
        
        self.text_area.clear()
        self.text_area.setText(display_text)
        
        scrollbar = self.text_area.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())
        
        if final:
            # 盖棺定论：这一段正式并入历史记忆变量
            self.committed_text = display_text
            
            # 同时爽快地推入剪贴板
            clipboard = QApplication.clipboard()
            clipboard.setText(self.committed_text)
            self.text_area.append("\n\n✅ [已自动复制草稿本的全部内容进剪贴板]")
            scrollbar.setValue(scrollbar.maximum())
            subprocess.Popen(["afplay", "/System/Library/Sounds/Glass.aiff"])

    def clear_notebook(self):
        """提供一个显式清空按钮，只有主人允许时才能擦除"""
        self.committed_text = ""
        self.text_area.clear()
        self.status_label.setText("🧹 草稿本已清空")
        self.status_label.setStyleSheet("font-size: 12px; color: orange; font-weight: bold;")

    # ---------------- 导出长文带时间戳功能 ----------------
    def import_audio_action(self):
        # 选择文件
        file_name, _ = QFileDialog.getOpenFileName(
            self, "选择你要转录的会议或长音频 (支持 .m4a, .mp3, .wav, .mp4)", 
            os.path.expanduser("~/Desktop"),
            "Audio/Video Files (*.m4a *.mp3 *.wav *.mp4 *.flac)"
        )
        if not file_name:
            return
            
        # 禁用录音功能，防止 GPU 打架
        self.record_btn.setEnabled(False)
        self.import_btn.setEnabled(False)
        self.status_label.setText("🚀 正在动用 GPU 疯狂转录中 (请勿关闭窗口)...")
        self.status_label.setStyleSheet("font-size: 12px; color: red; font-weight: bold;")
        
        # 将耗时的引擎任务扔进独立线程
        threading.Thread(target=self.run_export_task, args=(file_name,), daemon=True).start()

    def run_export_task(self, file_path):
        try:
            with self.mlx_lock:
                result = mlx_whisper.transcribe(
                    file_path,
                    path_or_hf_repo=MODEL,
                    language="zh"
                )
            
            segments = result.get('segments', [])
            
            if not segments:
                self.signals.export_status.emit("⚠️ 错误阶段", "没有识别到任何文字内容。")
                return

            # 构建 Word 文档
            doc = docx.Document()
            doc.add_heading('Whisper 本地转录报告', 0)
            
            original_filename = os.path.basename(file_path)
            doc.add_paragraph(f"源文件: {original_filename}")
            doc.add_paragraph(f"创建时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("-" * 40)

            for seg in segments:
                start_time = seg.get('start', 0.0)
                # 将 123.5 秒格式化为 [HH:MM:SS]
                m, s = divmod(start_time, 60)
                h, m = divmod(m, 60)
                time_str = f"[{int(h):02d}:{int(m):02d}:{int(s):02d}]"
                
                text_content = seg.get('text', '').strip()
                
                p = doc.add_paragraph()
                run_time = p.add_run(time_str + " ")
                run_time.bold = True
                run_time.font.color.rgb = docx.shared.RGBColor(128, 128, 128) # 灰色时间戳
                
                p.add_run(text_content)

            # 保存到桌面
            desktop = os.path.expanduser("~/Desktop")
            doc_name = f"转录_{original_filename}.docx"
            save_path = os.path.join(desktop, doc_name)
            doc.save(save_path)
            
            self.signals.export_status.emit(f"🎉 导出成功！Word 文件已保存在您的系统桌面。", "green")
            subprocess.Popen(["afplay", "/System/Library/Sounds/Glass.aiff"])
            
        except Exception as e:
            self.signals.export_status.emit(f"❌ 导出失败: {str(e)}", "red")

    def on_export_status(self, msg, color):
        # 恢复界面按钮状态并播报结果
        self.record_btn.setEnabled(True)
        self.import_btn.setEnabled(True)
        self.status_label.setText(msg)
        self.status_label.setStyleSheet(f"font-size: 12px; color: {color}; font-weight: bold;")
        
        if "成功" in msg:
            QMessageBox.information(self, "长录音转录完成", msg)
        else:
            QMessageBox.critical(self, "长录音转录异常", msg)

    def on_status_update(self, text, color):
        self.status_label.setText(text)
        self.status_label.setStyleSheet(f"font-size: 12px; color: {color}; font-weight: bold;")
        
    def on_btn_update(self, text, enabled):
        self.record_btn.setText(text)
        self.record_btn.setEnabled(enabled)

    def closeEvent(self, event):
        self.is_recording = False
        event.accept()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = VoiceInputApp()
    window.show()
    sys.exit(app.exec())
